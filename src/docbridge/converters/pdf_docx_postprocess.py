from __future__ import annotations

import logging
from pathlib import Path

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from lxml import etree

from docbridge.converters.md_docx_fonts import apply_docx_math_fallback_font, apply_docx_run_font
from docbridge.converters.md_theme import apply_a4_margins_2cm

logger = logging.getLogger(__name__)

_W_DRAWING = qn("w:drawing")

_WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
_WP_ANCHOR = f"{{{_WP_NS}}}anchor"
_WP_INLINE = f"{{{_WP_NS}}}inline"

_SKIP_ANCHOR_CHILD = frozenset(
    {
        "simplePos",
        "positionH",
        "positionV",
        "wrapNone",
        "wrapTopAndBottom",
        "wrapSquare",
        "wrapTight",
        "wrapThrough",
    }
)

_INLINE_CHILD_ORDER = {"extent": 0, "effectExtent": 1, "docPr": 2, "cNvGraphicFramePr": 3, "graphic": 4}


def _paragraph_has_drawing(p_element) -> bool:
    for node in p_element.iter():
        if node.tag == _W_DRAWING:
            return True
    return False


def _inline_child_sort_key(el) -> int:
    tag = etree.QName(el).localname
    return _INLINE_CHILD_ORDER.get(tag, 99)


def _iter_document_parts(doc: Document):
    yield doc.part
    seen = {id(doc.part)}
    for sec in doc.sections:
        for part in (getattr(sec.header, "part", None), getattr(sec.footer, "part", None)):
            if part is not None and id(part) not in seen:
                seen.add(id(part))
                yield part


def _convert_one_anchor_to_inline(anchor) -> None:
    parent = anchor.getparent()
    if parent is None:
        return

    to_move: list = []
    for child in list(anchor):
        tag = etree.QName(child).localname
        if tag in _SKIP_ANCHOR_CHILD:
            continue
        to_move.append(child)
    to_move.sort(key=_inline_child_sort_key)
    if not to_move:
        logger.warning("Skipping empty anchor (no positionable children)")
        return

    nsmap = anchor.nsmap
    if not nsmap and parent is not None:
        nsmap = parent.nsmap

    inline = etree.Element(_WP_INLINE, nsmap=nsmap)
    for k in ("distT", "distB", "distL", "distR"):
        inline.set(k, "0")

    for el in to_move:
        anchor.remove(el)
        inline.append(el)

    parent.remove(anchor)
    parent.append(inline)


def _convert_all_anchors_to_inline(doc: Document) -> int:
    n = 0
    for part in _iter_document_parts(doc):
        root = part.element
        anchors = [el for el in root.iter(_WP_ANCHOR)]
        for anchor in anchors:
            try:
                _convert_one_anchor_to_inline(anchor)
                n += 1
            except Exception as e:
                logger.warning("anchor→inline failed (image skipped): %s", e)
    if n:
        logger.info("Converted %d wp:anchor to wp:inline", n)
    return n


def _iter_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
    for sec in doc.sections:
        for part in (sec.header, sec.footer):
            if part is None:
                continue
            for p in part.paragraphs:
                yield p


def _apply_cjk_autospace_off(paragraph) -> None:
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    if p_pr.find(qn("w:autoSpaceDE")) is None:
        p_pr.insert(0, parse_xml(r'<w:autoSpaceDE {} w:val="0"/>'.format(nsdecls("w"))))
    if p_pr.find(qn("w:autoSpaceDN")) is None:
        p_pr.insert(0, parse_xml(r'<w:autoSpaceDN {} w:val="0"/>'.format(nsdecls("w"))))


def _font_name_is_tex_math_outline_face(name: str | None) -> bool:
    """
    仅「数学专用轮廓字」（非 LM Roman 正文）。用于映射到 Cambria Math，勿误伤 LMRoman10 等英文题名。
    """
    if not name:
        return False
    n = name.replace(" ", "").lower()
    if "cambria" in n:
        return False
    if "latinmodernmath" in n or "lmmath" in n:
        return True
    if any(k in n for k in ("cmsy", "cmmi", "cmex", "msam", "msbm")):
        return True
    if any(k in n for k in ("xitsmath", "stixmath", "notosansmath")):
        return True
    return False


def _font_name_suggests_math_or_technical(name: str | None) -> bool:
    """
    PDF（尤其 LaTeX）公式常用数学/符号字体；若再强制改为微软雅黑会大量缺字（□）。
    保留 pdf2docx 写入的字体名时一般能正常显示。
    """
    if not name:
        return False
    n = name.lower()
    return any(
        k in n
        for k in (
            "math",
            "cambria",
            "latin modern",
            "lmmath",
            "lmroman",
            "lm roman",
            "lmsans",
            "cmsy",
            "cmr",
            "cmmi",
            "cmex",
            "cmss",
            "xits",
            "stix",
            "dejavu",
            "symbol",
            "ams",
            "computer modern",
            "minion",
            "euler",
            "noto sans math",
            "libertine",
            "times new roman",
        )
    )


def _unicode_suggests_mathematics(text: str) -> bool:
    """行内碎片里若含数学专用 Unicode，禁止套正文无衬线中文字体。"""
    for ch in text:
        o = ord(ch)
        if (
            (0x0370 <= o <= 0x03FF)  # Greek
            or (0x1D400 <= o <= 0x1D7FF)  # Math alphanumeric symbols
            or (0x2100 <= o <= 0x214F)  # Letterlike symbols
            or (0x2190 <= o <= 0x23FF)  # Arrows + operators + misc technical
            or (0x27C0 <= o <= 0x27FF)  # Supplemental arrows, etc.
            or (0x2070 <= o <= 0x209F)  # Superscripts / subscripts
            or (0x2080 <= o <= 0x2089)  # Subscript digits
        ):
            return True
        if ch in "∫∑∏√∞≈≠≤≥±×·∂∇∀∃∩∪⊂⊃⊆⊇⊥⊤∧∨¬":
            return True
    return False


def _has_private_use_area_char(text: str) -> bool:
    """PDF 内嵌字体常用私用区码位；强制改字体易导致字形映射丢失。"""
    for ch in text:
        o = ord(ch)
        if 0xE000 <= o <= 0xF8FF:
            return True
        if 0xF0000 <= o <= 0xFFFFD or 0x100000 <= o <= 0x10FFFD:
            return True
    return False


def _needs_microsoft_yahei_body_font(text: str) -> bool:
    """
    后处理里只有「含中日韩表意/全角 CJK 标点」的片段才适合统一为微软雅黑正文。
    纯拉丁/希腊/数字/运算符的碎片多为 LaTeX 公式拆开后的 run，强改雅黑会大量 □。
    """
    for ch in text:
        o = ord(ch)
        if 0x4E00 <= o <= 0x9FFF:  # CJK Unified Ideographs
            return True
        if 0x3400 <= o <= 0x4DBF:  # Extension A
            return True
        if 0x3000 <= o <= 0x303F:  # CJK 标点与符号（、。《》等）
            return True
        if 0xFF00 <= o <= 0xFFEF:  # 全角字母数字与标点
            return True
        if 0xF900 <= o <= 0xFAFF:  # Compatibility ideographs
            return True
        if 0x2E80 <= o <= 0x2EFF:  # CJK Radicals Supplement / Kangxi
            return True
    return False


def _normalize_text_runs(paragraph) -> None:
    for run in paragraph.runs:
        if not run.text:
            continue
        fn = run.font.name or ""
        if "Courier" in fn or "Consolas" in fn or "Mono" in fn:
            continue
        if _font_name_suggests_math_or_technical(fn) or _unicode_suggests_mathematics(run.text):
            continue
        if _has_private_use_area_char(run.text):
            continue
        if not _needs_microsoft_yahei_body_font(run.text):
            continue
        sz = run.font.size
        bold = run.bold
        italic = run.italic
        apply_docx_run_font(run)
        if sz is not None:
            run.font.size = sz
        if bold is not None:
            run.bold = bold
        if italic is not None:
            run.italic = italic


def _remap_tex_math_fonts_to_cambria(paragraph) -> None:
    """未安装 TeX 数学字体时，Word/WPS 对 Latin Modern Math 常显示 □。"""
    for run in paragraph.runs:
        if not run.text:
            continue
        fn = run.font.name or ""
        if not _font_name_is_tex_math_outline_face(fn):
            continue
        sz = run.font.size
        bold = run.bold
        italic = run.italic
        apply_docx_math_fallback_font(run)
        if sz is not None:
            run.font.size = sz
        if bold is not None:
            run.bold = bold
        if italic is not None:
            run.italic = italic


def _strip_leading_empty_paragraphs(doc: Document) -> int:
    body = doc.element.body
    removed = 0
    while len(body) > 0:
        el = body[0]
        if el.tag != qn("w:p"):
            break
        p = Paragraph(el, doc)
        if p.text.strip() or _paragraph_has_drawing(el):
            break
        body.remove(el)
        removed += 1
    if removed:
        logger.info("Removed %d leading empty paragraphs", removed)
    return removed


def _clear_first_paragraph_space_before(doc: Document) -> None:
    body = doc.element.body
    for el in body:
        if el.tag != qn("w:p"):
            continue
        p = Paragraph(el, doc)
        if p.text.strip() or _paragraph_has_drawing(el):
            p.paragraph_format.space_before = Pt(0)
            break


def _patch_floating_anchors_allow_overlap(doc: Document) -> int:
    n = 0
    for part in _iter_document_parts(doc):
        for el in part.element.iter(_WP_ANCHOR):
            el.set("allowOverlap", "0")
            n += 1
    if n:
        logger.info("Set allowOverlap=0 on %d floating anchors", n)
    return n


def postprocess_pdf_docx(
    path: Path,
    *,
    normalize_fonts: bool = True,
    remap_tex_math_fonts: bool = True,
    match_margins: bool = False,
    cjk_autospace_fix: bool = True,
    convert_float_images_to_inline: bool = True,
    patch_floating_anchors: bool = False,
    trim_leading_empty_paragraphs: bool = True,
    clear_first_paragraph_space_before: bool = True,
) -> None:
    path = path.resolve()
    doc = Document(str(path))

    if match_margins:
        apply_a4_margins_2cm(doc)

    for p in _iter_paragraphs(doc):
        if cjk_autospace_fix:
            _apply_cjk_autospace_off(p)
        if normalize_fonts:
            _normalize_text_runs(p)
        if remap_tex_math_fonts:
            _remap_tex_math_fonts_to_cambria(p)

    if convert_float_images_to_inline:
        _convert_all_anchors_to_inline(doc)
    elif patch_floating_anchors:
        _patch_floating_anchors_allow_overlap(doc)

    if trim_leading_empty_paragraphs:
        _strip_leading_empty_paragraphs(doc)
    if clear_first_paragraph_space_before:
        _clear_first_paragraph_space_before(doc)

    doc.save(str(path))
