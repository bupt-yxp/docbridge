from __future__ import annotations

import logging
from pathlib import Path

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Inches, Pt, RGBColor

from docbridge.converters.md_common import resolve_resource_path
from docbridge.converters.md_docx_fonts import apply_docx_code_font, apply_docx_run_font
from docbridge.converters.md_theme import (
    BODY_PT,
    CODE_PT,
    HEADING_PT,
    LINE_HEIGHT,
    TEXT_RGB,
    apply_a4_margins_2cm,
    image_display_size_inches_safe,
)

logger = logging.getLogger(__name__)

_TEXT = RGBColor(*TEXT_RGB)


def html_fragment_to_docx(html_fragment: str, base_dir: Path, document: Document | None = None) -> Document:
    doc = document or Document()
    apply_a4_margins_2cm(doc)

    wrapped = f"<html><body>{html_fragment}</body></html>"
    soup = BeautifulSoup(wrapped, "html.parser")
    body = soup.body
    if not body:
        return doc

    for child in body.children:
        if isinstance(child, NavigableString) and not str(child).strip():
            continue
        _element_to_docx(child, doc, base_dir, list_level=0)

    return doc


def _add_theme_heading(doc: Document, text: str, level: int) -> None:
    level = min(max(level, 1), 6)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level <= 2 else 6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = LINE_HEIGHT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(HEADING_PT[level])
    run.font.color.rgb = _TEXT
    apply_docx_run_font(run)


def _finalize_body_runs(paragraph) -> None:
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = LINE_HEIGHT
    for r in paragraph.runs:
        if not r.text:
            continue
        fn = r.font.name or ""
        if "Courier" in fn or "Mono" in fn or "Consolas" in fn:
            r.font.size = Pt(CODE_PT)
            r.font.color.rgb = _TEXT
            apply_docx_code_font(r)
            continue
        r.font.size = Pt(BODY_PT)
        r.font.color.rgb = _TEXT
        apply_docx_run_font(r)


def _element_to_docx(node: Tag | NavigableString, doc: Document, base_dir: Path, list_level: int) -> None:
    if isinstance(node, NavigableString):
        return
    if not isinstance(node, Tag):
        return

    name = node.name.lower()

    if name in {"h1", "h2", "h3", "h4", "h5", "h6"}:
        level = min(int(name[1]), 9)
        _add_theme_heading(doc, node.get_text(strip=True), level)
        return

    if name == "p":
        p = doc.add_paragraph()
        _add_inline_runs(p, node, base_dir)
        _finalize_body_runs(p)
        return

    if name == "blockquote":
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        _add_inline_runs(p, node, base_dir)
        _finalize_body_runs(p)
        return

    if name in {"ul", "ol"}:
        _add_list(node, doc, base_dir, list_level, ordered=name == "ol")
        return

    if name == "table":
        _add_table(node, doc, base_dir)
        return

    if name == "pre":
        code = node.find("code")
        text = code.get_text() if code else node.get_text()
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = LINE_HEIGHT
        run = p.add_run(text)
        run.font.size = Pt(CODE_PT)
        run.font.color.rgb = _TEXT
        apply_docx_code_font(run)
        return

    if name == "hr":
        p = doc.add_paragraph("—" * 20)
        _finalize_body_runs(p)
        return

    if name in {"div", "section"}:
        for sub in node.children:
            _element_to_docx(sub, doc, base_dir, list_level)
        return

    if name not in {"br"}:
        p = doc.add_paragraph()
        _add_inline_runs(p, node, base_dir)
        _finalize_body_runs(p)


def _add_list(node: Tag, doc: Document, base_dir: Path, list_level: int, ordered: bool) -> None:
    style = "List Number" if ordered else "List Bullet"
    for li in node.find_all("li", recursive=False):
        if not isinstance(li, Tag):
            continue
        p = doc.add_paragraph(style=style)
        if list_level:
            p.paragraph_format.left_indent = Pt(22 * list_level)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = LINE_HEIGHT
        for child in li.children:
            if isinstance(child, NavigableString):
                t = str(child)
                if t.strip():
                    r = p.add_run(t)
                    r.font.size = Pt(BODY_PT)
                    r.font.color.rgb = _TEXT
                    apply_docx_run_font(r)
            elif isinstance(child, Tag):
                cname = child.name.lower()
                if cname in {"ul", "ol"}:
                    _add_list(child, doc, base_dir, list_level + 1, ordered=cname == "ol")
                elif cname == "p":
                    _add_inline_runs(p, child, base_dir)
                else:
                    _add_inline_runs(p, child, base_dir)
        for r in p.runs:
            if not r.text:
                continue
            fn = r.font.name or ""
            if "Courier" in fn or "Mono" in fn or "Consolas" in fn:
                r.font.size = Pt(CODE_PT)
                apply_docx_code_font(r)
                continue
            r.font.size = Pt(BODY_PT)
            r.font.color.rgb = _TEXT
            apply_docx_run_font(r)


def _add_table(table: Tag, doc: Document, base_dir: Path) -> None:
    rows = table.find_all("tr")
    if not rows:
        return
    col_counts = [len(tr.find_all(["th", "td"])) for tr in rows]
    ncols = max(col_counts) if col_counts else 1
    nrows = len(rows)
    t = doc.add_table(rows=nrows, cols=ncols)
    t.style = "Table Grid"
    for ri, tr in enumerate(rows):
        cells = tr.find_all(["th", "td"])
        for ci, cell in enumerate(cells):
            if ci >= ncols:
                break
            tc = t.rows[ri].cells[ci]
            tc.text = ""
            inner_ps = cell.find_all("p")
            if inner_ps:
                for pi, pp in enumerate(inner_ps):
                    para = tc.paragraphs[0] if pi == 0 else tc.add_paragraph()
                    _add_inline_runs(para, pp, base_dir)
                    _finalize_body_runs(para)
            else:
                _add_inline_runs(tc.paragraphs[0], cell, base_dir)
                _finalize_body_runs(tc.paragraphs[0])
            if cell.name == "th":
                for para in tc.paragraphs:
                    for r in para.runs:
                        r.bold = True
                        r.font.color.rgb = _TEXT
                        if r.text:
                            apply_docx_run_font(r)


def _add_inline_runs(paragraph, parent: Tag, base_dir: Path, bold: bool = False, italic: bool = False) -> None:
    for child in parent.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if text:
                r = paragraph.add_run(text)
                r.bold, r.italic = bold, italic
            continue
        if not isinstance(child, Tag):
            continue
        tag = child.name.lower()

        if tag in {"strong", "b"}:
            _add_inline_runs(paragraph, child, base_dir, True, italic)
        elif tag in {"em", "i"}:
            _add_inline_runs(paragraph, child, base_dir, bold, True)
        elif tag == "code":
            r = paragraph.add_run(child.get_text())
            r.font.size = Pt(CODE_PT)
            r.bold, r.italic = bold, italic
            r.font.color.rgb = _TEXT
            apply_docx_code_font(r)
        elif tag == "a":
            r = paragraph.add_run(child.get_text())
            r.bold, r.italic = bold, italic
        elif tag == "br":
            paragraph.add_run().add_break()
        elif tag == "img":
            src = child.get("src")
            path = resolve_resource_path(base_dir, src)
            if path:
                try:
                    w_in, h_in = image_display_size_inches_safe(path)
                    paragraph.add_run().add_picture(str(path), width=Inches(w_in), height=Inches(h_in))
                except Exception as e:
                    logger.warning("Image insert failed %s: %s", path, e)
                    r = paragraph.add_run(f"[image: {src}]")
                    r.font.size = Pt(BODY_PT)
                    r.font.color.rgb = _TEXT
                    apply_docx_run_font(r)
            else:
                r = paragraph.add_run(f"[image: {src}]")
                r.font.size = Pt(BODY_PT)
                r.font.color.rgb = _TEXT
                apply_docx_run_font(r)
        else:
            _add_inline_runs(paragraph, child, base_dir, bold, italic)
