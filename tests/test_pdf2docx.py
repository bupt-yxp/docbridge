from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document as DocxDocument

from docbridge.base import ConversionOptions
from docbridge.converters.pdf_docx import PdfToDocxConverter
from docbridge.converters.pdf_docx_postprocess import (
    _font_name_is_tex_math_outline_face,
    _font_name_suggests_math_or_technical,
    _has_private_use_area_char,
    _needs_microsoft_yahei_body_font,
    _unicode_suggests_mathematics,
)


def test_pdf_to_docx_delegates_to_sibling_md_for_omml_math(tmp_path: Path) -> None:
    """与 PDF 同名的 .md 存在时改走 MD→DOCX，公式为 Word OMML，而非 pdf2docx 纯文字。"""
    fitz = pytest.importorskip("fitz")
    md = tmp_path / "doc.md"
    md.write_text("# T\n\nInline $x^2$.\n", encoding="utf-8")
    pdf = tmp_path / "doc.pdf"
    pdf_doc = fitz.open()
    page = pdf_doc.new_page()
    page.insert_text((72, 72), "x", fontsize=12)
    pdf_doc.save(pdf)
    pdf_doc.close()

    out = tmp_path / "out.docx"
    PdfToDocxConverter().convert(
        pdf,
        out,
        ConversionOptions(md_backend="python", pdf_auto_use_sibling_markdown=True),
    )
    assert out.is_file()
    ddoc = DocxDocument(str(out))
    assert "m:oMath" in ddoc.part.element.xml


def test_pdf_table_parse_defaults_disabled_for_latex_formula_pdfs() -> None:
    o = ConversionOptions()
    assert o.pdf_parse_lattice_table is False
    assert o.pdf_parse_stream_table is False


def test_pdf_postprocess_skips_yahei_for_math_heuristics() -> None:
    assert _font_name_suggests_math_or_technical("Latin Modern Math")
    assert _font_name_suggests_math_or_technical("LM Roman 10")
    assert not _font_name_suggests_math_or_technical("Microsoft YaHei")
    assert _unicode_suggests_mathematics("∫₀¹ f(x)dx")
    assert _unicode_suggests_mathematics("α + β = π")
    assert not _unicode_suggests_mathematics("仅中文与ASCII")
    assert not _needs_microsoft_yahei_body_font("ax^2+bx+c")
    assert not _needs_microsoft_yahei_body_font("lim")
    assert _needs_microsoft_yahei_body_font("二次方程")
    assert _needs_microsoft_yahei_body_font("，")  # 全角逗号：仍走雅黑统一
    assert _has_private_use_area_char("\uE000")
    assert _font_name_is_tex_math_outline_face("Latin Modern Math")
    assert not _font_name_is_tex_math_outline_face("LM Roman 10")


def test_minimal_pdf_to_docx(tmp_path: Path) -> None:
    fitz = pytest.importorskip("fitz")
    pdf = tmp_path / "minimal.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "DocBridge test", fontsize=12)
    doc.save(pdf)
    doc.close()

    out = tmp_path / "out.docx"
    PdfToDocxConverter().convert(pdf, out, ConversionOptions(render_dpi=200.0))
    assert out.is_file()
    assert out.stat().st_size > 500


def test_landscape_pdf_preserves_orientation_after_postprocess(tmp_path: Path) -> None:
    fitz = pytest.importorskip("fitz")
    pdf = tmp_path / "landscape.pdf"
    doc = fitz.open()
    page = doc.new_page(width=842, height=595)
    page.insert_text((72, 72), "landscape", fontsize=12)
    doc.save(pdf)
    doc.close()

    out = tmp_path / "out.docx"
    PdfToDocxConverter().convert(
        pdf,
        out,
        ConversionOptions(render_dpi=200.0, pdf_postprocess=True, pdf_match_page_margins=False),
    )
    assert out.is_file()
    d = DocxDocument(str(out))
    sec = d.sections[0]
    assert sec.page_width > sec.page_height


def test_pdf_match_md_theme_a4_forces_portrait(tmp_path: Path) -> None:
    fitz = pytest.importorskip("fitz")
    pdf = tmp_path / "landscape.pdf"
    doc = fitz.open()
    page = doc.new_page(width=842, height=595)
    page.insert_text((72, 72), "landscape", fontsize=12)
    doc.save(pdf)
    doc.close()

    out = tmp_path / "out_a4.docx"
    PdfToDocxConverter().convert(
        pdf,
        out,
        ConversionOptions(
            render_dpi=200.0,
            pdf_postprocess=True,
            pdf_match_page_margins=True,
        ),
    )
    d = DocxDocument(str(out))
    sec = d.sections[0]
    assert sec.page_height > sec.page_width
